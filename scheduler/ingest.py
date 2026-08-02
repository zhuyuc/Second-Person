"""
文档导入 Ingest（产品文档 §外部素材 Ingest / 开发文档 §四）。

- 支持 PDF / DOCX / TXT / MD / URL
- 原始文件存 raw_docs/（不可变）+ raw_docs 表元数据
- 长文档按语义边界切分为不超过 ingest_chunk_tokens 的块，块间 200 token 重叠
- 每块独立走 Distiller 提炼与归属判定；silent_doc_import 决定静默或预览
- 单文件上限 50 MB；raw_docs 总量超 2 GB 告警
"""
from __future__ import annotations

import asyncio
import logging
from datetime import datetime
from pathlib import Path

from memory.naming import raw_doc_id
from infrastructure.timeutil import now_cst

logger = logging.getLogger("second_person.ingest")

MAX_FILE_MB = 50
RAW_TOTAL_WARN_GB = 2
OVERLAP_CHARS = 500  # 约 200 token
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif"}


def image_to_data_url(path: Path) -> str:
    """把本地图片读为 OpenAI 多模态所需的 base64 dataURL。"""
    import base64
    import mimetypes
    mime = mimetypes.guess_type(str(path))[0] or "image/png"
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{b64}"


async def ocr_extract_text(path: Path) -> str:
    """本地 OCR 引擎（RapidOCR，纯 onnxruntime，无系统二进制依赖）。

    未安装或识别失败时返回空串，由调用方降级处理。仅提取文字，不理解图表语义。
    """
    try:
        from rapidocr_onnxruntime import RapidOCR  # 延迟导入，未装则降级
    except Exception:  # noqa: BLE001
        logger.info("未安装 rapidocr_onnxruntime，OCR 引擎不可用")
        return ""
    try:
        engine = RapidOCR()
        # RapidOCR 为同步计算，放线程池避免阻塞事件循环
        result, _ = await asyncio.to_thread(engine, str(path))
        if not result:
            return ""
        return "\n".join(line[1] for line in result if len(line) > 1).strip()
    except Exception as e:  # noqa: BLE001
        logger.warning("OCR 解析失败：%s", e)
        return ""


async def extract_text_async(path: Path, image_fn=None, mime: str = "") -> str:
    """统一异步文本提取入口：图片走 image_fn（VLM/OCR），其余走同步 extract_text。

    PDF/DOCX 走富解析：文字 + 内嵌图片（逐张过 image_fn）+ 表格（DOCX 转 Markdown）。
    image_fn 为空或返回空时，图片按“仅缓存不解析”降级（返回空串）。
    """
    ext = path.suffix.lower()
    if ext in IMAGE_EXTS:
        if image_fn is None:
            return ""
        return (await image_fn(path)) or ""
    if ext == ".docx":
        try:
            return await _extract_docx_rich(path, image_fn)
        except Exception as e:  # noqa: BLE001
            logger.warning("DOCX 富解析失败，降级纯文字提取：%s", e)
            return await asyncio.to_thread(extract_text, path, mime)
    if ext == ".pdf":
        try:
            return await _extract_pdf_rich(path, image_fn)
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF 富解析失败，降级纯文字提取：%s", e)
            return await asyncio.to_thread(extract_text, path, mime)
    # 纯文本/其余格式同样丢线程：大文件读盘解码也不占事件循环
    return await asyncio.to_thread(extract_text, path, mime)


# ---- 混合文档富解析：文字 + 内嵌图片 + 表格 ------------------------------
_MIN_IMAGE_BYTES = 5 * 1024   # 小于 5KB 的图（图标/装饰线）跳过
_MIN_IMAGE_SIZE = 64          # 宽或高小于 64px 跳过


async def _image_blob_to_text(blob: bytes, suffix: str, image_fn) -> str:
    """内嵌图片字节 → 临时文件 → image_fn 解析。噪声图（过小）返回空。"""
    if not image_fn or not blob or len(blob) < _MIN_IMAGE_BYTES:
        return ""
    try:
        import io
        from PIL import Image
        with Image.open(io.BytesIO(blob)) as im:
            if im.width < _MIN_IMAGE_SIZE or im.height < _MIN_IMAGE_SIZE:
                return ""
    except Exception:  # noqa: BLE001
        pass  # 尺寸检查失败不阻断，继续尝试解析
    import tempfile
    suffix = suffix if suffix.lower() in IMAGE_EXTS else ".png"
    tmp = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
            f.write(blob)
            tmp = Path(f.name)
        return (await image_fn(tmp)) or ""
    except Exception as e:  # noqa: BLE001
        logger.warning("内嵌图片解析失败（已跳过）：%s", e)
        return ""
    finally:
        if tmp is not None:
            try:
                tmp.unlink()
            except OSError:
                pass


def _docx_table_to_md(table) -> str:
    """DOCX 表格转 Markdown 管道表；合并单元格取左上值、重复引用置空。"""
    lines = []
    for r in table.rows:
        seen_tc = set()
        cells = []
        for c in r.cells:
            key = id(c._tc)
            if key in seen_tc:
                cells.append("")
            else:
                seen_tc.add(key)
                cells.append(" ".join(c.text.split()))
        lines.append("| " + " | ".join(cells) + " |")
    if not lines:
        return ""
    ncols = len(table.rows[0].cells)
    sep = "| " + " | ".join(["---"] * ncols) + " |"
    return "\n".join([lines[0], sep] + lines[1:])


def _para_image_rids(para_el) -> list[str]:
    """提取段落内嵌图片的关系 ID（a:blip 的 r:embed）。"""
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main",
          "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    return [blip.get(f"{{{ns['r']}}}embed")
            for blip in para_el.findall(".//a:blip", ns)
            if blip.get(f"{{{ns['r']}}}embed")]


def _docx_collect(path: Path) -> list[tuple]:
    """同步解析 DOCX 结构（CPU 密集，须在工作线程执行）。

    按文档流原始顺序返回 [("text", 内容) | ("img", 字节, 后缀)] 有序列表。
    """
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    d = docx.Document(str(path))
    items: list[tuple] = []
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:tbl"):
            md = _docx_table_to_md(Table(child, d))
            if md:
                items.append(("text", md))
        elif child.tag == qn("w:p"):
            para = Paragraph(child, d)
            if para.text.strip():
                items.append(("text", para.text))
            # 段落内嵌图片：解析结果插回原位置，保持上下文语义
            for rid in _para_image_rids(child):
                part = d.part.related_parts.get(rid)
                if part is None:
                    continue
                suffix = Path(getattr(part, "partname", "x.png")
                              ).suffix or ".png"
                items.append(("img", part.blob, suffix))
    return items


async def _extract_docx_rich(path: Path, image_fn) -> str:
    """DOCX 富解析：段落 / 表格(Markdown) / 内嵌图片(VLM)。

    结构解析丢工作线程，事件循环只做图片 VLM/OCR 的异步调用。
    """
    items = await asyncio.to_thread(_docx_collect, path)
    parts: list[str] = []
    for item in items:
        if item[0] == "text":
            parts.append(item[1])
        else:
            text = await _image_blob_to_text(item[1], item[2], image_fn)
            if text:
                parts.append(f"【图片内容】{text}")
    return "\n".join(parts)


async def _extract_pdf_rich(path: Path, image_fn) -> str:
    """PDF 富解析：逐页文字 + 该页内嵌图片（pypdf 提字节 → VLM）追加页尾。

    pdfplumber/pypdf 解析均为 CPU 密集，丢工作线程执行。
    """
    def _pdf_texts() -> list[str]:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return [p.extract_text() or "" for p in pdf.pages]

    page_texts = await asyncio.to_thread(_pdf_texts)
    # 无解析回调时与旧行为一致，只返回文字
    if image_fn is None:
        return "\n".join(page_texts)

    def _pdf_image_blobs() -> list[tuple]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        out: list[tuple] = []
        for i, page in enumerate(reader.pages):
            for img in page.images:
                suffix = Path(img.name or "x.png").suffix or ".png"
                out.append((i, img.data, suffix))
        return out

    page_images: dict[int, list[str]] = {}
    try:
        blobs = await asyncio.to_thread(_pdf_image_blobs)
        for i, data, suffix in blobs:
            text = await _image_blob_to_text(data, suffix, image_fn)
            if text:
                page_images.setdefault(i, []).append(f"【图片内容】{text}")
    except ImportError:
        logger.info("未安装 pypdf，PDF 内嵌图片跳过解析")
    except Exception as e:  # noqa: BLE001
        logger.warning("PDF 内嵌图片提取失败（仅保留文字）：%s", e)
    parts: list[str] = []
    for i, text in enumerate(page_texts):
        if text:
            parts.append(text)
        parts.extend(page_images.get(i, []))
    return "\n".join(parts)


def extract_text(path: Path, mime: str = "") -> str:
    ext = path.suffix.lower()
    # 常见纯文本/代码/数据格式：直接解码
    text_exts = {
        ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".log", ".yaml", ".yml",
        ".xml", ".html", ".htm", ".ini", ".conf", ".toml",
        ".py", ".js", ".ts", ".jsx", ".tsx", ".vue", ".java", ".go", ".rs", ".c",
        ".cpp", ".h", ".hpp", ".cs", ".rb", ".php", ".sh", ".sql", ".css", ".scss",
    }
    if ext in text_exts:
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages)
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF 解析失败：%s", e)
            return ""
    if ext == ".docx":
        try:
            import docx
            d = docx.Document(str(path))
            return "\n".join(p.text for p in d.paragraphs)
        except Exception as e:  # noqa: BLE001
            logger.warning("DOCX 解析失败：%s", e)
            return ""
    # 其余格式：尝试 UTF-8 解码（容错），二进制无法解析则返回空
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:  # noqa: BLE001
        return ""


def chunk_text(text: str, chunk_chars: int) -> list[str]:
    """按段落/空行边界切分，块间保留重叠。

    兜底两种退化情形，避免长文档被当作单个 chunk（导致 ingest_chunk_tokens 失效）：
    - 正文无空行（如 DOCX 段落以单换行拼接）→ 回退按单换行分段；
    - 单个段落本身超长 → 按 chunk_chars 硬切。
    """
    if len(text) <= chunk_chars:
        return [text] if text.strip() else []
    paras = text.split("\n\n")
    if len(paras) == 1:
        paras = text.split("\n")
    # 超长段落硬切，保证单段不会撑爆一个 chunk
    units: list[str] = []
    for p in paras:
        if len(p) > chunk_chars:
            units.extend(p[i:i + chunk_chars]
                         for i in range(0, len(p), chunk_chars))
        else:
            units.append(p)
    chunks, buf, size = [], [], 0
    for p in units:
        buf.append(p)
        size += len(p)
        if size >= chunk_chars:
            chunk = "\n\n".join(buf)
            chunks.append(chunk)
            # 重叠
            tail = chunk[-OVERLAP_CHARS:]
            buf, size = [tail], len(tail)
    if buf and "".join(buf).strip():
        chunks.append("\n\n".join(buf))
    return chunks


class IngestManager:
    def __init__(self, db, data_dir, distiller, config, notifier=None,
                 image_extract_fn=None):
        self.db = db
        self.data_dir = Path(data_dir)
        self.distiller = distiller
        self.config = config
        self.notify = notifier or (lambda t, m: None)
        # 图片解析回调（VLM/OCR/off 分派，由 container 装配）；缺省则图片仅缓存不解析
        self.image_extract_fn = image_extract_fn
        self.raw_dir = self.data_dir / "raw_docs"
        self.raw_dir.mkdir(parents=True, exist_ok=True)
        # 容量告警标志：超阈值首次告警，恢复后重置，避免每次导入重复推送
        self._capacity_alerted = False

    def _next_doc_id(self) -> str:
        # 以现存 id 的最大数字后缀 +1 生成，而非 count(*)+1：
        # 删除文档会在序号中留下空洞，count 方案会重算出已存在的较小 id
        # 从而触发 raw_docs.id UNIQUE 约束冲突，导致后续导入全部失败。
        row = self.db.query_one(
            "SELECT MAX(CAST(SUBSTR(id, 5) AS INTEGER)) m FROM raw_docs "
            "WHERE id LIKE 'doc_%'")
        seq = (row["m"] or 0) + 1
        # 收尾兜底：极端并发/历史脏数据下仍可能撞号，递增直至空闲
        while self.db.query_one(
                "SELECT 1 FROM raw_docs WHERE id=?", (raw_doc_id(seq),)):
            seq += 1
        return raw_doc_id(seq)

    async def ingest_file(self, filename: str, content: bytes,
                          source: str = "web_ui", progress_cb=None) -> dict:
        """progress_cb(stage, data)（可选）：流式导入时向前端推送实时进度，
        stage 依次 extracting/chunked/distilling；为 None 时与旧同步行为一致。"""
        size_mb = len(content) / (1024 * 1024)
        if size_mb > MAX_FILE_MB:
            raise ValueError(f"文件超过 {MAX_FILE_MB} MB 上限")
        doc_id = self._next_doc_id()
        stored = self.raw_dir / f"{doc_id}_{filename}"
        stored.write_bytes(content)
        # 落盘与提炼、入库非原子：任一环节失败则回滚删除已落盘文件，
        # 避免留下无数据库记录的孤儿文件（列表读 DB，孤儿文件永不可见/不可管理）。
        preview_mode = (not self.config.get("silent_doc_import", True)
                        and source == "web_ui")
        try:
            if progress_cb:
                await progress_cb("extracting", {})
            text = await extract_text_async(stored, self.image_extract_fn)
            is_image = stored.suffix.lower() in IMAGE_EXTS
            chunk_chars = self.config.get(
                "ingest_chunk_tokens", 6000) * 2  # token→char 粗略
            chunks = chunk_text(text, chunk_chars)
            if progress_cb:
                await progress_cb("chunked", {"total": len(chunks)})

            if preview_mode:
                # 预览模式：只提炼+归属判定不写入，暂存待用户勾选确认
                items: list[dict] = []
                failed = 0
                last_err: Exception | None = None
                for idx, chunk in enumerate(chunks):
                    if progress_cb:
                        await progress_cb(
                            "distilling", {"current": idx + 1, "total": len(chunks)})
                    try:
                        items.extend(await self.distiller.distill_preview(
                            chunk, source_type="knowledge"))
                    except Exception as e:  # noqa: BLE001
                        failed += 1
                        last_err = e
                if chunks and failed == len(chunks):
                    raise last_err or RuntimeError("文档提炼全部失败")
                import json
                now = now_cst().isoformat(timespec="seconds")
                self.db.execute(
                    "INSERT INTO raw_docs(id,filename,file_path,file_size,mime_type,"
                    "source,extracted_memory_ids,imported_at,extracted_text,"
                    "review_status) VALUES(?,?,?,?,?,?,?,?,?,'pending')",
                    (doc_id, filename, str(stored.relative_to(self.data_dir)),
                     len(content), stored.suffix, source, "[]", now,
                     text if is_image else None))
                self.db.execute(
                    "INSERT OR REPLACE INTO pending_imports(doc_id,items,created_at) "
                    "VALUES(?,?,?)",
                    (doc_id, json.dumps(items, ensure_ascii=False), now))
                self._check_capacity()
                return {"doc_id": doc_id, "preview": True,
                        "items": [{"index": i, "title": it.get("title", ""),
                                   "summary": it.get("summary", ""),
                                   "attribution": it.get("attribution", "imported"),
                                   "confidence": it.get("confidence", "medium")}
                                  for i, it in enumerate(items)]}

            written: list[str] = []
            failed = 0
            last_err: Exception | None = None
            for idx, chunk in enumerate(chunks):
                if progress_cb:
                    await progress_cb(
                        "distilling", {"current": idx + 1, "total": len(chunks)})
                # 单块提炼失败不中断整篇，但记录失败数用于整体判定
                try:
                    ids = await self.distiller.distill(chunk, source_type="knowledge")
                    written.extend(ids)
                except Exception as e:  # noqa: BLE001
                    failed += 1
                    last_err = e
                    logger.warning("文档 %s 第 %d 块提炼失败：%s",
                                   doc_id, failed, e)
            # 所有分块均失败（如模型熔断）→ 视为导入失败，回滚并抛出原始错误
            if chunks and failed == len(chunks):
                raise last_err or RuntimeError("文档提炼全部失败")
        except (Exception, asyncio.CancelledError):
            # 含前端断开（刷新/关页）导致的任务取消：CancelledError 继承自
            # BaseException，必须显式捕获才能回滚落盘文件，否则残留孤儿文件。
            try:
                stored.unlink()
            except OSError:
                pass
            raise

        import json
        self.db.execute(
            "INSERT INTO raw_docs(id,filename,file_path,file_size,mime_type,source,"
            "extracted_memory_ids,imported_at,extracted_text) VALUES(?,?,?,?,?,?,?,?,?)",
            (doc_id, filename, str(stored.relative_to(self.data_dir)), len(content),
             stored.suffix, source, json.dumps(written, ensure_ascii=False),
             now_cst().isoformat(timespec="seconds"),
             text if is_image else None))
        self._check_capacity()
        # 图片启用解析后仍无文本：明确告知用户仅缓存，消除“以为已收录”的静默隐患
        if is_image and not (text or "").strip():
            self.notify("image_not_parsed",
                        f"图片「{filename}」未解析（未启用视觉/OCR 或模型不支持），仅缓存")
        if self.config.get("silent_doc_import", True):
            msg = f"已从文档提取 {len(written)} 条记忆"
            if failed:
                msg += f"（{failed} 个分块提炼失败已跳过）"
            self.notify("doc_imported", msg)
        result = {"doc_id": doc_id, "extracted": len(written),
                  "memory_ids": written}
        if failed:
            result["partial_failed"] = failed
        return result

    async def ingest_url(self, url: str, web_fetch_fn) -> dict:
        text = await web_fetch_fn(url)
        doc_id = self._next_doc_id()
        import json
        chunk_chars = self.config.get("ingest_chunk_tokens", 6000) * 2
        written: list[str] = []
        for chunk in chunk_text(text, chunk_chars):
            written.extend(await self.distiller.distill(chunk, source_type="knowledge"))
        self.db.execute(
            "INSERT INTO raw_docs(id,filename,file_path,file_size,mime_type,source,"
            "source_url,extracted_memory_ids,imported_at) VALUES(?,?,?,?,?,?,?,?,?)",
            (doc_id, url[:80], "", len(text), "text/html", "url", url,
             json.dumps(written, ensure_ascii=False),
             now_cst().isoformat(timespec="seconds")))
        self.notify("doc_imported", f"已从网页提取 {len(written)} 条记忆")
        return {"doc_id": doc_id, "extracted": len(written)}

    async def confirm_import(self, doc_id: str, selected: list[int]) -> dict:
        """预览确认：按用户勾选的下标写入记忆，未勾选的丢弃。
        完成后 raw_docs 标 reviewed，回顾 Agent 不再重提炼该文档（尊重用户筛选）。"""
        import json
        row = self.db.query_one(
            "SELECT items FROM pending_imports WHERE doc_id=?", (doc_id,))
        if not row:
            raise KeyError(f"无待确认导入：{doc_id}")
        items = json.loads(row["items"] or "[]")
        selected_set = {int(i) for i in (selected or [])}
        written: list[str] = []
        for i, item in enumerate(items):
            if i not in selected_set:
                continue
            try:
                mid = await self.distiller.write_item(item, source_type="knowledge")
                if mid:
                    written.append(mid)
            except Exception:  # noqa: BLE001
                logger.warning("预览确认写入失败：doc=%s idx=%d", doc_id, i)
        self.db.execute(
            "UPDATE raw_docs SET extracted_memory_ids=?, review_status='reviewed' "
            "WHERE id=?", (json.dumps(written, ensure_ascii=False), doc_id))
        self.db.execute(
            "DELETE FROM pending_imports WHERE doc_id=?", (doc_id,))
        self.notify("doc_imported",
                    f"已按筛选写入 {len(written)} 条记忆（丢弃 {len(items) - len(written)} 条）")
        return {"doc_id": doc_id, "written": len(written),
                "discarded": len(items) - len(written), "memory_ids": written}

    def list_documents(self) -> list[dict]:
        import json
        rows = self.db.query_all(
            "SELECT * FROM raw_docs ORDER BY imported_at DESC")
        return [{"id": r["id"], "filename": r["filename"], "size": r["file_size"],
                 "imported_at": r["imported_at"],
                 "memory_count": len(json.loads(r["extracted_memory_ids"] or "[]"))}
                for r in rows]

    def get_document_detail(self, doc_id: str) -> dict | None:
        """文档详情：元数据 + 重新解析的正文 + 提炼出的记忆列表。"""
        import json
        row = self.db.query_one(
            "SELECT * FROM raw_docs WHERE id=?", (doc_id,))
        if not row:
            return None
        mem_ids = json.loads(row["extracted_memory_ids"] or "[]")
        memories: list[dict] = []
        if mem_ids:
            ph = ",".join("?" * len(mem_ids))
            mrows = self.db.query_all(
                f"SELECT id,title,summary FROM memories WHERE id IN ({ph})", mem_ids)
            order = {mid: i for i, mid in enumerate(mem_ids)}
            mrows = sorted(mrows, key=lambda r: order.get(r["id"], 999))
            memories = [{"id": r["id"], "title": r["title"],
                         "summary": r["summary"]} for r in mrows]
        content = ""
        if row["file_path"]:
            f = self.data_dir / row["file_path"]
            if f.exists():
                content = extract_text(f)
        # 被引用记录：该文档提炼的记忆被对话引用的明细（知识库侧使用凭证）
        cites = self.db.query_all(
            "SELECT ce.memory_id, ce.session_id, ce.cited_at, "
            "s.title AS session_title, m.title AS memory_title "
            "FROM citation_events ce "
            "LEFT JOIN sessions s ON ce.session_id=s.session_id "
            "LEFT JOIN memories m ON ce.memory_id=m.id "
            "WHERE ce.doc_id=? ORDER BY ce.cited_at DESC LIMIT 50", (doc_id,))
        return {
            "id": row["id"], "filename": row["filename"], "size": row["file_size"],
            "mime_type": row["mime_type"], "source": row["source"],
            "source_url": row["source_url"], "imported_at": row["imported_at"],
            "content": content, "memories": memories,
            "memory_count": len(mem_ids),
            "citations": [{"memory_id": r["memory_id"],
                           "memory_title": r["memory_title"] or r["memory_id"],
                           "session_id": r["session_id"],
                           "session_title": r["session_title"] or r["session_id"],
                           "cited_at": r["cited_at"]} for r in cites],
        }

    def delete_document(self, doc_id: str) -> None:
        row = self.db.query_one(
            "SELECT file_path FROM raw_docs WHERE id=?", (doc_id,))
        if row and row["file_path"]:
            f = self.data_dir / row["file_path"]
            if f.exists():
                f.unlink()
        self.db.execute("DELETE FROM raw_docs WHERE id=?", (doc_id,))
        self.db.execute(
            "DELETE FROM pending_imports WHERE doc_id=?", (doc_id,))

    def _check_capacity(self) -> None:
        total = sum(f.stat().st_size for f in self.raw_dir.rglob(
            "*") if f.is_file())
        over = total > RAW_TOTAL_WARN_GB * 1024 ** 3
        if over and not self._capacity_alerted:
            # 首次超阈值才告警：容量告警是持续状态，每次导入重复推送会刷屏
            self._capacity_alerted = True
            self.notify("raw_docs_capacity",
                        f"raw_docs 总大小超过 {RAW_TOTAL_WARN_GB} GB，建议清理")
        elif not over:
            self._capacity_alerted = False  # 恢复后重置，下次超阈值重新告警

    def cleanup_temp_attachments(self, days: int = 7) -> int:
        """清理 data/temp/attachments/ 下超 N 天的临时文件（夜间维护链）。"""
        import time
        temp_dir = self.data_dir / "temp" / "attachments"
        if not temp_dir.exists():
            return 0
        cutoff = time.time() - days * 86400
        removed = 0
        for f in temp_dir.rglob("*"):
            if f.is_file() and f.stat().st_mtime < cutoff:
                try:
                    f.unlink()
                    removed += 1
                except OSError:
                    pass
        return removed
